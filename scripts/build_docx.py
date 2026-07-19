#!/usr/bin/env python3
"""Render a minutes-spec JSON into a .docx meeting-minutes document.

Usage:
    python scripts/build_docx.py <minutesspec.json> [--out deliverables/<name>.docx]

Template resolution: meta.template in templates/ -> first templates/*.docx ->
built-in clean style. See specs/minutes-spec.schema.md.

Page budget: python-docx cannot paginate, so the >10-page check is an estimate
(word count / 450 words-per-page). Estimate over budget -> exit code 2 so the
calling agent compresses and re-renders.
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INK_2 = RGBColor(0x52, 0x51, 0x4E)
MUTED = RGBColor(0x89, 0x87, 0x81)
ACCENT = RGBColor(0x2A, 0x78, 0xD6)

WORDS_PER_PAGE = 450
PAGE_BUDGET = 10


def resolve_template(meta, templates_dir):
    name = meta.get("template")
    if name:
        p = templates_dir / name
        if p.exists():
            return p
        print(f"WARN: template '{name}' not found; falling back", file=sys.stderr)
    for p in sorted(templates_dir.glob("*.docx")):
        return p
    return None


def new_document(meta, templates_dir):
    tpl = resolve_template(meta, templates_dir)
    if tpl:
        doc = Document(str(tpl))
        # Template supplies styles/header/footer; clear its body content.
        for el in list(doc.element.body):
            if el.tag.endswith("}sectPr"):
                continue
            doc.element.body.remove(el)
        print(f"Using template: {tpl.name}")
    else:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        print("Using built-in clean style")
    return doc


def add_kv_line(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.color.rgb = INK_2
    r2 = p.add_run(str(value))
    r2.font.color.rgb = INK_2


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("spec")
    ap.add_argument("--out", default=None)
    ap.add_argument("--templates", default="templates")
    args = ap.parse_args()

    spec = json.loads(Path(args.spec).read_text(encoding="utf-8"))
    meta = spec.get("meta", {})
    doc = new_document(meta, Path(args.templates))

    # Title block
    h = doc.add_heading(meta.get("title", "Meeting Minutes"), level=0)
    for kv in (("Date", meta.get("date")), ("Time", meta.get("time")),
               ("Location", meta.get("location")), ("Author", meta.get("author"))):
        if kv[1]:
            add_kv_line(doc, *kv)

    # Attendees
    attendees = spec.get("attendees", [])
    if attendees:
        doc.add_heading("Attendees", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1" if "Light Grid Accent 1" in [s.name for s in doc.styles] else table.style
        hdr = table.rows[0].cells
        for i, t in enumerate(("Name", "Role", "Present")):
            hdr[i].text = t
        for a in attendees:
            row = table.add_row().cells
            row[0].text = a.get("name", "")
            row[1].text = a.get("role", "")
            row[2].text = "Yes" if a.get("present", True) else "No"

    # Agenda
    agenda = spec.get("agenda", [])
    if agenda:
        doc.add_heading("Agenda", level=1)
        for i, item in enumerate(agenda, 1):
            doc.add_paragraph(f"{i}. {item}")

    # Sections
    for sec in spec.get("sections", []):
        doc.add_heading(sec["heading"], level=1)
        summary = sec.get("summary")
        if summary:
            p = doc.add_paragraph()
            r = p.add_run(summary)
            r.italic = True
            r.font.color.rgb = INK_2
        for pt in sec.get("points", []):
            text = pt["text"]
            if pt.get("confidence") == "L":
                text += " (L)"
            doc.add_paragraph(text, style="List Bullet")

    # Decisions
    decisions = spec.get("decisions", [])
    if decisions:
        doc.add_heading("Decisions", level=1)
        for d in decisions:
            line = d["text"] + (f' — owner: {d["owner"]}' if d.get("owner") else "")
            doc.add_paragraph(line, style="List Bullet")

    # Actions
    actions = spec.get("actions", [])
    if actions:
        doc.add_heading("Action Items", level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        for i, t in enumerate(("Action", "Owner", "Due")):
            hdr[i].text = t
        for a in actions:
            row = table.add_row().cells
            row[0].text = a.get("text", "")
            row[1].text = a.get("owner", "")
            row[2].text = a.get("due", "TBD")

    # Human review box
    review = spec.get("review_items", [])
    if review:
        doc.add_heading("Needs Human Review", level=1)
        for r_item in review:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f'{r_item["text"]} — {r_item.get("reason", "low confidence")}')
            run.font.color.rgb = ACCENT

    # Sources appendix
    sources = spec.get("appendix_sources", [])
    if sources:
        doc.add_heading("Appendix — Sources", level=1)
        for s in sources:
            p = doc.add_paragraph(s, style="List Bullet")
            for run in p.runs:
                run.font.color.rgb = MUTED
                run.font.size = Pt(9)

    out = args.out
    if not out:
        safe = re.sub(r"[^\w\- ]", "", meta.get("title", "minutes")).strip().replace(" ", "-")
        out = f"deliverables/{safe or 'minutes'}.docx"
    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    words = sum(len(p.text.split()) for p in doc.paragraphs)
    words += sum(len(c.text.split()) for t in doc.tables for row in t.rows for c in row.cells)
    est_pages = max(1, round(words / WORDS_PER_PAGE + 0.5))
    print(f"Saved {out} (~{est_pages} pages estimated, {words} words)")
    if est_pages > PAGE_BUDGET:
        print(f"ERROR: estimated {est_pages} pages > {PAGE_BUDGET}-page budget. Compress and re-render.",
              file=sys.stderr)
        sys.exit(2)


if __name__ == "__main__":
    main()
